"""File parsing service — PDF / DOCX / TXT / MD / JSON → plain text."""

from __future__ import annotations

import io
import json
import logging
from typing import Optional

logger = logging.getLogger(__name__)

_MAX_TEXT_CHARS = 150_000
_PDF_MIN_CHARS = 50

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "application/json",
}


class FileParseError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def detect_format(content: str, mime_type: str) -> str:
    if mime_type == "application/json":
        return "json"
    if mime_type == "text/markdown" or content.strip().startswith("#"):
        return "markdown"
    return "text"


def _truncate_at_sentence(text: str, limit: int) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    # Truncate at nearest sentence boundary below limit
    truncated = text[:limit]
    for sep in (". ", ".\n", "! ", "!\n", "? ", "?\n"):
        idx = truncated.rfind(sep)
        if idx != -1:
            truncated = truncated[: idx + 1]
            return truncated, True
    # Fall back to last space
    idx = truncated.rfind(" ")
    if idx != -1:
        truncated = truncated[:idx]
    return truncated, True


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FileParseError("internal_error", "pypdf not available") from exc

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    text = "\n".join(parts).strip()
    if len(text) < _PDF_MIN_CHARS:
        raise FileParseError(
            "scanned_pdf",
            f"PDF yielded only {len(text)} chars — likely a scanned image without embedded text",
        )
    return text


def _parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise FileParseError("internal_error", "python-docx not available") from exc

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise FileParseError("file_parse_error", f"DOCX could not be opened: {exc}") from exc

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def _parse_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise FileParseError("file_parse_error", "File is not valid UTF-8 text") from exc


def _parse_json(data: bytes) -> str:
    try:
        raw = data.decode("utf-8")
        obj = json.loads(raw)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except UnicodeDecodeError as exc:
        raise FileParseError("file_parse_error", "JSON file is not valid UTF-8") from exc
    except json.JSONDecodeError as exc:
        raise FileParseError("file_parse_error", f"Invalid JSON: {exc}") from exc


def parse_uploaded_file(data: bytes, mime_type: str, filename: str) -> tuple[str, str, bool]:
    """Parse a file and return (parsed_text, format, truncated).

    Raises FileParseError on unsupported type or parse failure.
    """
    if mime_type not in SUPPORTED_MIME_TYPES:
        raise FileParseError(
            "unsupported_file_type",
            f"File type '{mime_type}' is not supported. Allowed: PDF, DOCX, TXT, MD, JSON",
        )

    if mime_type == "application/pdf":
        text = _parse_pdf(data)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = _parse_docx(data)
    elif mime_type == "application/json":
        text = _parse_json(data)
    else:
        # text/plain and text/markdown
        text = _parse_text(data)

    text, truncated = _truncate_at_sentence(text, _MAX_TEXT_CHARS)
    fmt = detect_format(text, mime_type)
    return text, fmt, truncated
